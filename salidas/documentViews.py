from django.shortcuts import render, render_to_response, RequestContext
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from django.http import HttpResponse
from django.contrib import auth
from django.shortcuts import render,render_to_response,redirect,get_object_or_404
from django.contrib.auth import  authenticate
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required,user_passes_test,permission_required
from django.core.urlresolvers import reverse
# from salidas.forms import *
from salidas.models import Application, Document
from datetime import date
import locale
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from io import StringIO
from docx import * #to generate Docs
import io,os,os.path,tempfile, zipfile

locale.setlocale(locale.LC_ALL, '')

def getfiles(request):
    print("get files")
    # Files (local path) to put in the .zip
    filenames=[]
    id_app = request.GET['id']
    app = Application.objects.get(pk = id_app)
    filenames.append(solicitud_facultad_doc(app))


    solicitante=str(app.id_Teacher)
    replacements=app.get_replacements()
    replacement_teachers=set()

    for replacement in replacements:
        replacement_teachers.add(replacement.rut_teacher)
    if len(replacement_teachers)==1:
        filenames.append(carta_reemplazo_ac_doc(app,replacement))
    else:
        for replacement in replacements:
            filenames.append(carta_reemplazo_any(app,replacement))

    filenames.append(peticion_docente_doc(app,replacement_teachers))

    # Folder name in ZIP archive which contains the above files
    # E.g [thearchive.zip]/somefiles/file2.txt
    zip_subdir = "documentos"
    zip_filename = "%s.zip" % zip_subdir

    # Open StringIO to grab in-memory ZIP contents
    f = io.BytesIO()
    # The zip compressor
    zf = zipfile.ZipFile(f, "w")

    for fpath in filenames:
        # Calculate path for file in zip
        fdir, fname = os.path.split(fpath)
        zip_path = os.path.join(zip_subdir, fname)

        # Add file, at correct path
        zf.write(fpath, zip_path)

    # Must close zip for all contents to be written
    zf.close()

    # Grab ZIP file from in-memory, make response with correct MIME-type
    resp = HttpResponse(f.getvalue(),content_type='application/zip')
    # ..and correct content-disposition
    resp['Content-Disposition'] = 'attachment; filename=%s' % zip_filename

    return resp

def peticion_docente_doc(app,replacement_teachers):
    replacements=app.get_replacements()
    print("peticion docente")
    path = os.path.join(settings.MEDIA_ROOT, 'carta_peticion_docente.docx') #todo:path para cada profe (?)

    if len(replacement_teachers)==1:
        replacement_type="En mis actividades académicas y docentes seré reemplazado por"+str(replacement_teachers.pop())
    else:
        i=0
        replace=[]
        for rep in replacements:
            if(str(rep.type)=="Docente"):
                replace.append("docentes")
            else:
                replace.append("académicas")
            i+=1
            replace.append(str(rep.rut_teacher))
        replacement_type="En mis actividades " +replace[0]+ " seré reemplazado por " +replace[1]+ " y en mis actividades "+ replace[2]+ " seré reemplazado por " +replace[3]+"."


    # user = request.user.id
    document = Document()
    document.add_paragraph("Eric Tanter").alignment = 2 #
    document.add_paragraph("DCC - UChile").alignment = 2
    document.add_paragraph("etanter@dcc.uchile.cl").alignment = 2
    p = document.add_paragraph()
    p.add_run().add_break()
    p.add_run('Señor ').bold = True
    p.add_run().add_break()
    p.add_run('Sergio Ochoa').bold = True
    p.add_run().add_break()
    p.add_run('Director DCC').bold = True
    p.add_run().add_break()
    p.add_run('Presente').bold = True
    p.add_run().add_break()

    reemplazo=document.add_paragraph(replacement_type)

    document.save(path)
    return path

def solicitud_facultad_doc(app):
    print("solicitud fac")
    document = Document()
    title = document.add_paragraph('SOLICITUD DE COMISION O PERMISO')
    title.alignment = 1
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break()
    p.add_run('UNIDAD ACADEMICA ').bold = True
    p.add_run('     DEPARTAMENTO DE CIENCIAS DE LA COMPUTACION')
    path = os.path.join(settings.MEDIA_ROOT, 'solicitud_facultad.doc')
    document.save(path)
    return path

def carta_reemplazo_any(app,replacement):
    solicitante=str(app.id_Teacher)
    replacement_teacher=replacement.rut_teacher
    replacement_type=replacement.type
    return carta_reemplazo(solicitante,app,replacement_teacher,replacement_type)

def carta_reemplazo_ac_doc(app,replacement_teacher):
    solicitante=str(app.id_Teacher)
    replacement_type="académico y docente"

    return carta_reemplazo(solicitante,app,replacement_teacher,replacement_type)

def carta_reemplazo(solicitante,app,replacement_teacher,replacement_type):

    signature_file=replacement_teacher.signature
    try:
        signature_path=str(signature_file)
        #print('firma')
        #print(signature_path)
        signature_path=signature_path.split("/")
        #print(signature_path)
        signature_path=signature_path[1]
        #print(signature_path)
    except:
        print('no hay firma')
    fecha_inicio=app.get_start_date().strftime("%A %d %B %Y")
    fecha_fin=app.get_end_date().strftime("%A %d %B %Y")

    print("carta de reemplazo")
    filename= 'compromiso_reemplazo'+str(replacement_type)+'.doc'
    path = os.path.join(settings.MEDIA_ROOT,filename)
    document=Document()

    today=date.today().strftime("%A %d %B %Y")
    today_date=document.add_paragraph(str(today))
    today_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    today_date.add_run().add_break()

    title = document.add_paragraph('CARTA DE COMPROMISO DE REEMPLAZO')
    title.alignment = 1
    title.add_run().add_break()
    title.add_run().add_break()
    title.add_run().add_break()
    cuerpo='Yo, '+str(replacement_teacher)+', me comprometo a reemplazar al señor académico '+solicitante+', del '+fecha_inicio+ ' al '+fecha_fin+', ambas fechas inclusive, en todas sus actividades del tipo '+str(replacement_type)+'.'
    jerarquia=replacement_teacher.hierarchy
    jornada=replacement_teacher.working_day

    paragraph = document.add_paragraph(cuerpo)
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()
    paragraph.add_run().add_break()

    p = document.add_paragraph()
    p.add_run().add_break()
    p.add_run('Nombre: ').bold = True
    p.add_run(str(replacement_teacher))
    p.add_run().add_break()
    p.add_run('Jerarquía: ').bold = True
    p.add_run(str(jerarquia))
    p.add_run().add_break()
    p.add_run('Cargo: ').bold = True
    p.add_run('Académico Jornada ')
    p.add_run(str(jornada))
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run().add_break()
    p.add_run('                                                                                                               ')
    try:
        p.add_run().add_picture(os.path.join(settings.MEDIA_ROOT,"signatures",signature_path),width=Inches(1.0)).alignment = WD_ALIGN_PARAGRAPH.RIGHT#signature_file.path)
    except:
        print("no hay firma")

    signature=document.add_paragraph()
    signature.add_run('                                                                  '+str(replacement_teacher)).bold=True
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature.add_run().add_break()
    signature.add_run().add_break()
    signature.add_run().add_break()

    other_signatures=document.add_paragraph()
    other_signatures.add_run('V°B° Sr. Director').bold = True
    other_signatures.add_run().add_break()
    other_signatures.add_run().add_break()
    other_signatures.add_run('V°B° Jefa Docente').bold = True
    #document.add_paragraph('first item in unordered list', style='ListBullet')
    #document.add_heading('The role of dolphins', level=2)
    #document.add_picture('image-filename.png')
    document.save(path)
    return path